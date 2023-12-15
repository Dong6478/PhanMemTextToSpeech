from pydub import AudioSegment
import math
import os
from docx import Document
import openai
import json
import tkinter as tk
from tkinter import filedialog
from tkinter import simpledialog, messagebox, scrolledtext
import asyncio
import queue
import threading
import time
from EncryptPass import ma_hoa, giai_ma 

# Tạo queue để truyền dữ liệu giữa các luồng
data_queue = queue.Queue()
transcript_content = ""  # Biến toàn cục để lưu trữ nội dung

# Hàm để cắt tệp âm thanh thành các đoạn
async def split_audio(input_file, output_folder, max_length=125):
    audio = AudioSegment.from_file(input_file, format="wav")
    num_segments = math.ceil(len(audio) / (max_length * 1000))
    file_name = os.path.splitext(os.path.basename(input_file))[0]

    doc = Document()

    for i in range(num_segments):
        start_time = i * max_length * 1000
        end_time = (i + 1) * max_length * 1000
        segment = audio[start_time:end_time]
        output_file = os.path.join(output_folder, f"{file_name}_segment_{i+1}.wav")
        output_file = output_file.replace("\\","/")
        segment.export(output_file, format="wav")

        try:
            text = await audio_to_text_and_translate(output_file)
            data_queue.put(text)  # Đưa dữ liệu vào queue
            doc.add_paragraph(text)
             
            # Đổ dữ liệu từ danh sách vào ScrolledText
            fill_scrolled_text(transcript_text, text)
            print(text + "\n")
            
        except Exception as e:
            print(f"Lỗi khi chuyển đổi tệp: {output_file}, lỗi: {e}")
            break

        delete_generated_files(output_file)

    word_file = os.path.join(output_folder, f"{file_name}_translated.docx")
    word_file = word_file.replace("\\","/")
    doc.save(word_file)
    return word_file


# Hàm để lấy API key từ file hoặc từ người dùng
def get_api_key():
    try:
        with open("api_key.txt", "r") as file:
            return file.read().strip()
    except FileNotFoundError:
        return None

# Hàm để lưu API key mới
def set_api_key(new_key):
    with open("api_key.txt", "w") as file:
        file.write(new_key)

# Hàm để yêu cầu API key mới từ người dùng
def request_new_api_key():
    root = tk.Tk()
    root.withdraw()
    new_key = simpledialog.askstring("API Key Hết Hạn", "Vui lòng nhập API key mới:", parent=root)
    if new_key:
        set_api_key(new_key)
        messagebox.showinfo("Thông Báo", "Thay key API mới thành công")
        return new_key
    
    return None

# Hàm để chuyển đổi âm thanh thành văn bản
async def audio_to_text_and_translate(audio_file_path):
    api_key = get_api_key()
    if (api_key == ""):
        api_key = "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"

    with open(audio_file_path, "rb") as audio_file:
        openai.api_key = api_key
        try:
            transcript = openai.Audio.transcribe("whisper-1", audio_file)
            parsed_data = json.loads(f"{transcript}")
            text = parsed_data['text']
            await asyncio.sleep(21)
            return text
        except openai.error.OpenAIError as e:
            messagebox.showerror("Lỗi API", f"Lỗi API: {e}")
            new_key = request_new_api_key()
            return None

# Hàm xóa tệp tạo ra
def delete_generated_files(file_path):
    os.remove(file_path)

# Hàm xử lý tất cả tệp trong thư mục
async def process_directory(input_dir, output_dir):
    for file_name in os.listdir(input_dir):
        if file_name.endswith('.wav'):
            input_file = os.path.join(input_dir, file_name).replace("\\","/")
            try:
                message = f"Đọc nội dung file video:  {file_name}"
                fill_scrolled_text(transcript_text, message)
                word_file = await split_audio(input_file, output_dir)
                print(f"Tệp Word đã tạo: {word_file}")
            except Exception as e:
                print(f"Lỗi khi xử lý tệp: {input_file}, lỗi: {e}")
                break

# Biến toàn cục để lưu trữ đường dẫn thư mục
input_directory = ''
output_directory = ''

# Hàm thực hiện xử lý
async def process():
    if input_directory and output_directory:
        transcript_text.delete('1.0', tk.END)
        transcript_text.insert(tk.END, "Bắt đầu xử lý...\n")
        await process_directory(input_directory, output_directory)
        transcript_text.insert(tk.END, "Xử lý hoàn tất!")
    else:
        print("Vui lòng chọn cả thư mục Input và Output.")
        

# Thêm hàm cập nhật các hộp văn bản
def update_input_textbox():
    input_textbox.delete(0, tk.END)
    input_textbox.insert(0, input_directory)

def update_output_textbox():
    output_textbox.delete(0, tk.END)
    output_textbox.insert(0, output_directory)

# Sửa đổi các hàm select_input_folder và select_output_folder
def select_input_folder():
    global input_directory
    input_directory = filedialog.askdirectory()
    update_input_textbox()

def select_output_folder():
    global output_directory
    output_directory = filedialog.askdirectory()
    update_output_textbox()

# Hàm bắt đầu quá trình xử lý trong một luồng riêng biệt
# Sửa đổi hàm start_process để bắt đầu lưu file
def start_process():
    global output_directory
    if input_directory and output_directory:
        info_label.config(text="Đang xử lý...")
        save_thread = threading.Thread(target=save_text_to_file, args=(output_directory,))
        save_thread.daemon = True  # Đảm bảo luồng tự đóng khi chương trình kết thúc
        save_thread.start()

        process_thread = threading.Thread(target=lambda: asyncio.run(process()))
        process_thread.start()
    else:
        messagebox.showinfo("Thông Báo", "Vui lòng chọn cả thư mục Input và Output.")
        
def update_transcript_content():
    global transcript_content
    transcript_content = transcript_text.get("1.0", tk.END)
        
def fill_scrolled_text(text_widget, item):
    text_widget.insert(tk.END, f"{item}\n")
    
    
# Hàm mới để lưu nội dung vào file .docx
def save_text_to_file(output_dir):
    global transcript_content
    while True:
        try:
            # Tạo đối tượng Document mới
            doc = Document()
            update_transcript_content()
            # Thêm nội dung vào document
            doc.add_paragraph(transcript_content)
            # Lưu file
            doc.save(os.path.join(output_dir, "FileTotalTextAll.docx"))
            time.sleep(0.2)  # Lưu file mỗi 100ms
        except:
            print("Đang mở file không thể lưu")
            messagebox.showerror("Lỗi Lưu File", "Đang mở file docx không thể lưu nội dung chỉnh sửa. Vui lòng đóng file docx để có thể lưu nội dung.")

def close_app():
    # Đặt logic để dừng tất cả các tiến trình và luồng đang chạy
    # Ví dụ: dừng luồng, đóng kết nối, v.v.
    root = tk.Tk()
    root.quit()  # Dừng vòng lặp sự kiện của Tkinter
    root.destroy()  # Đóng cửa sổ ứng dụng


# Hàm main để tạo GUI
def main():
    global input_textbox, output_textbox, info_label, transcript_text
    
    root = tk.Tk()
    # Gắn hàm đóng ứng dụng với sự kiện đóng cửa sổ
    root.protocol("WM_DELETE_WINDOW", close_app)
    root.title("Chọn Thư Mục và Xử Lý")
    root.configure(bg='lightblue')

    tk.Button(root, text="Chọn Thư Mục Input", command=select_input_folder, bg='lightgrey').pack(pady=5, fill='x', padx=200)
    input_textbox = tk.Entry(root, width=50, font=("Arial", 16))    
    input_textbox.pack(pady=5, padx=20, fill='x')

    tk.Button(root, text="Chọn Thư Mục Output", command=select_output_folder, bg='lightgrey').pack(pady=5, fill='x', padx=200)
    output_textbox = tk.Entry(root, width=50, font=("Arial", 16))
    output_textbox.pack(pady=5, padx=20, fill='x')

    tk.Button(root, text="Thực Hiện Xử Lý", command=start_process, bg='lightgreen').pack(pady=10)

    # Thêm label để hiển thị thông báo trạng thái
    info_label = tk.Label(root, text="", bg='lightblue')
    info_label.pack(pady=5)
    
     # Thêm Text widget để hiển thị nội dung
    transcript_text = scrolledtext.ScrolledText(root, wrap=tk.WORD)
    transcript_text.pack(pady=10, padx=10, expand=True, fill="both")
    
    # Thay đổi cỡ chữ cho transcript_text
    transcript_text.config(font=("Arial", 16))  # Thay đổi cỡ chữ tại đây
    
    root.mainloop()

if __name__ == "__main__":
    main()

